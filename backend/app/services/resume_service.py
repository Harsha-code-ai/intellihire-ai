"""
Resume parsing service.
Supports PDF (pdfplumber → pypdf fallback) and DOCX (python-docx).
Production-safe: every code path returns a string, never raises.
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger("intellihire.resume")

# Maximum characters we bother to keep (saves memory & DB space)
MAX_TEXT_LENGTH = 15_000


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from a resume file.

    Returns a non-empty string on success.
    Returns a fallback message string on any failure (never raises).
    """
    try:
        if not file_bytes:
            logger.warning("extract_text: empty file_bytes received")
            return ""

        ext = Path(filename).suffix.lower().strip()

        if ext == ".pdf":
            text = _extract_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            text = _extract_docx(file_bytes)
        elif ext in (".txt", ".text", ".md"):
            text = _extract_plain(file_bytes)
        else:
            # Unknown extension — try plain text first, then PDF
            logger.info(f"Unknown extension '{ext}' for '{filename}', trying plain-text then PDF")
            text = _extract_plain(file_bytes)
            if not text.strip():
                text = _extract_pdf(file_bytes)

        cleaned = _clean(text)
        if cleaned:
            return cleaned[:MAX_TEXT_LENGTH]

        logger.warning(f"No readable text found in '{filename}'")
        return ""

    except Exception as e:
        logger.error(f"extract_text crashed for '{filename}': {e}", exc_info=True)
        return ""


# ─── PDF ──────────────────────────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    """Try pdfplumber first, fall back to pypdf."""

    # ── Attempt 1: pdfplumber ────────────────────────────────────────────────
    try:
        import pdfplumber  # noqa: F401

        parts: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            if not pdf.pages:
                raise ValueError("PDF has no pages")
            for i, page in enumerate(pdf.pages):
                try:
                    t = page.extract_text()
                    if t and t.strip():
                        parts.append(t)
                except Exception as page_err:
                    logger.debug(f"pdfplumber page {i} error: {page_err}")
                    continue

        if parts:
            return "\n".join(parts)
        logger.debug("pdfplumber returned empty text; trying pypdf")

    except ImportError:
        logger.debug("pdfplumber not installed; trying pypdf")
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")

    # ── Attempt 2: pypdf ─────────────────────────────────────────────────────
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        parts = []
        for i, page in enumerate(reader.pages):
            try:
                t = page.extract_text()
                if t and t.strip():
                    parts.append(t)
            except Exception as page_err:
                logger.debug(f"pypdf page {i} error: {page_err}")
                continue

        if parts:
            return "\n".join(parts)
        logger.warning("pypdf also returned empty text")

    except ImportError:
        logger.error("pypdf not installed — cannot parse PDF")
    except Exception as e:
        logger.error(f"pypdf failed: {e}", exc_info=True)

    return ""


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())

        return "\n".join(lines)

    except ImportError:
        logger.error("python-docx not installed — cannot parse DOCX")
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}", exc_info=True)

    return ""


# ─── Plain text ───────────────────────────────────────────────────────────────

def _extract_plain(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    """Remove null bytes and excessive whitespace."""
    if not text:
        return ""
    text = text.replace("\x00", "")
    # Collapse runs of 3+ blank lines into 2
    import re
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()